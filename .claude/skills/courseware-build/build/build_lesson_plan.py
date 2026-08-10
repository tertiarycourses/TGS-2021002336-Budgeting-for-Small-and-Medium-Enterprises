#!/usr/bin/env python3
"""Generate the WSQ Budgeting for Small and Medium Enterprises Lesson Plan (LP)
DOCX in the Tertiary house format.

Cover page + Document Version Control Record + auto TOC + Arial 11pt body +
colour-coded 2-day schedule tables (9:30am-6:30pm, 8 training hours/day, 1h
lunch, tea within, final assessment Day 2 from 4:00pm). Topics/labs come from
course_data + the domain data files so the LP stays aligned with the deck,
guide and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5; from data_domain6 import DOMAIN6
from data_domain7 import DOMAIN7
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5+DOMAIN6+DOMAIN7
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
HEADER_FILL="1F6FEB"; TOPIC_FILL="E8F0FE"; BREAK_FILL="FFF4E5"; LUNCH_FILL="FDE9D9"; ASSESS_FILL="E8F7EE"

def lab_titles(nums):
    return "; ".join(f"Lab {a['num']}: {a['title']}" for a in ACT if a['num'] in nums)

# ------------------------------------------------ schedule (single source of truth for timing)
# (start, end, minutes, kind, activity_text)  kind: admin/topic/lab/break/lunch/assess/recap
SCHEDULE = {
 1: (C.DAY_THEMES[1], [
    ("9:30","10:00",30,"admin","Digital Attendance (AM). Welcome, trainer and learners introduction, ground rules, learning outcomes and course outline (Slides 1–17)"),
    ("10:00","11:00",60,"topic","Topic 1 — Introduction to Financial Budgeting: what is a budget, the three financial statements, budget goals, master budget, stakeholders and budget types (Slides 18–32)"),
    ("11:00","11:15",15,"break","Tea break"),
    ("11:15","13:00",105,"lab","Hands-on: "+lab_titles([1,2,3])+" (Slides 33–40)"),
    ("13:00","14:00",60,"lunch","Lunch break. Digital Attendance (PM)"),
    ("14:00","15:45",105,"topic","Topic 2 — Financial Forecasting: accounting principles, accrual vs cash, corporate finance, time value of money, forecasting techniques (Slides 41–51). Hands-on: "+lab_titles([4])+" (Slides 52–54)"),
    ("15:45","16:00",15,"break","Tea break"),
    ("16:00","18:15",135,"topic","Topic 3 — Budget Preparation: budget process, budgeting methods, cash budget, working capital, the 7-step master budget (Slides 55–67). Hands-on: "+lab_titles([5])+" (Slides 68–70)"),
    ("18:15","18:30",15,"recap","Day 1 recap and Q&A"),
 ]),
 2: (C.DAY_THEMES[2], [
    ("9:30","9:45",15,"admin","Digital Attendance (AM). Day 1 recap"),
    ("9:45","10:45",60,"topic","Topic 4 — Budget Control Plan: budgetary control loop, operating/cash/capex controls, road map (Slides 72–78). Hands-on: "+lab_titles([6])+" (Slides 79–81)"),
    ("10:45","11:00",15,"break","Tea break"),
    ("11:00","13:00",120,"topic","Topic 5 — Budget Analysis: budget vs actual vs forecast, variances, thresholds, analysis techniques, KPIs, technology and BI (Slides 82–92). Hands-on: "+lab_titles([7,8])+" (Slides 93–98)"),
    ("13:00","14:00",60,"lunch","Lunch break. Digital Attendance (PM)"),
    ("14:00","14:30",30,"topic","Topic 6 — Budget Approval: approval process, stakeholder precautions, common issues (Slides 100–104). Hands-on: "+lab_titles([9])+" (Slides 105–107)"),
    ("14:30","15:30",60,"topic","Topic 7 — Financial Compliance: Singapore corporate taxation, YA, start-up exemption, ECI and filing, penalties (Slides 108–114). Hands-on: "+lab_titles([10,11,12])+" (Slides 115–122)"),
    ("15:30","15:45",15,"break","Tea break"),
    ("15:45","16:00",15,"assess","Course feedback and TRAQOM survey. Briefing for Assessment (Slides 123–130). Digital Attendance (Assessment)"),
    ("16:00","17:10",70,"assess","Written Assessment (WA) — Short-Answer Questions (SAQ), 70 minutes, open book"),
    ("17:10","18:30",80,"assess","Written Assessment (WA) — Case Study (CS), 80 minutes, open book"),
 ]),
}

# ------------------------------------------------ build document
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc,"LESSON PLAN",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("10",C.VERSION_DATE,"Legacy release — Budgeting for SMEs 2-day lesson plan (reference deck v10).",C.TRAINER),
 ("11",C.VERSION_DATE,"Full regeneration in the single-source pipeline: 12 hands-on labs across 7 topics, visual slide deck, aligned LG and assessments (WA SAQ 70 min + WA CS 80 min).",C.TRAINER),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,"QA alignment pass: slide references re-verified against deck v12; break dividers corrected (End of Day 1, Day 2 lunch).",C.TRAINER),
])
prodoc.add_toc(doc)

def H(text,level=1):
    h=doc.add_heading(text,level=level); return h

H("Course Information",1)
info=[("Course Title",C.TITLE),("WSQ Course Reference",C.COURSE_CODE),
      ("Skills Framework TSC",f"{C.TSC_TITLE} ({C.TSC_CODE})"),
      ("Training Provider",C.ORG+"  ("+C.UEN.replace('UEN: ','UEN ')+")"),
      ("Duration","2 days · 8 training hours per day (16 hours)"),
      ("Daily Timing","9:30 am – 6:30 pm (1-hour lunch; tea breaks within training time)"),
      ("Mode","Instructor-led, hands-on budgeting activities on Xero, Power BI and case-study templates"),
      ("Trainer",C.TRAINER)]
t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
for k,v in info:
    c=t.add_row().cells; c[0].text=""; r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(10)
    prodoc._shade_cell(c[0],TOPIC_FILL)
    c[1].text=""; c[1].paragraphs[0].add_run(v).font.size=Pt(10)

H("Learning Outcomes",1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size=Pt(10.5)

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "Final assessment is conducted on Day 2 from 4:00 pm (SAQ), followed by the Case Study.",C.ASSESSMENT["note"]]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(10.5)

def set_cell(cell,text,bold=False,size=9.5,color=None,fill=None,align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"
    if color: r.font.color.rgb=color
    if fill: prodoc._shade_cell(cell,fill)

KIND_FILL={"topic":TOPIC_FILL,"break":BREAK_FILL,"lunch":LUNCH_FILL,"assess":ASSESS_FILL,
           "admin":"F3F5F8","recap":"F3F5F8","lab":None}

H("Course Schedule",1)
for day,(theme,rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}",2)
    tbl=doc.add_table(rows=0,cols=3); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.add_row().cells
    for i,htext in enumerate(["Time","Duration","Topic / Activity"]):
        set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
    training=0
    for start,end,mins,kind,text in rows:
        cells=tbl.add_row().cells; fill=KIND_FILL.get(kind)
        set_cell(cells[0],f"{start}–{end}",bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[1],f"{mins} min",size=9.5,fill=fill)
        set_cell(cells[2],text,bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        if kind!="lunch": training+=mins
    for row in tbl.rows:
        row.cells[0].width=Inches(1.15); row.cells[1].width=Inches(0.9); row.cells[2].width=Inches(4.75)
    p=doc.add_paragraph(); r=p.add_run(f"Total training time: {training} minutes ({training//60} hours)."); r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    assert training==480, f"Day {day} training minutes = {training}, expected 480"

H("Lab Reference (aligned to the Skills Framework TSC)",1)
tt=doc.add_table(rows=0,cols=3); tt.style="Table Grid"
hdr=tt.add_row().cells
for i,htext in enumerate(["Topic","TSC Coverage","Labs"]):
    set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
for tp in C.TOPICS:
    acts=[a for a in ACT if a["topic"]==tp["num"]]
    cells=tt.add_row().cells
    set_cell(cells[0],f"Topic {tp['code']}: {tp['title']}",bold=True,size=9.5,fill=TOPIC_FILL)
    set_cell(cells[1],tp["weighting"],size=9.5,fill=TOPIC_FILL)
    set_cell(cells[2],", ".join(f"Lab {a['num']}" for a in acts),size=9.5)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT=os.path.join(REPO,"courseware",f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved",OUT)
