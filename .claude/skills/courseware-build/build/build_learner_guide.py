#!/usr/bin/env python3
"""Generate the WSQ Budgeting for Small and Medium Enterprises Learner Guide as
BOTH a Markdown mirror (LG-*.md at repo root) and a DOCX (courseware/LG-*.docx)
from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per lab (Objective · Goal · What you'll produce · Detailed
step-by-step · Check your work), plus setup, revision and glossary. All content
is driven by course_data + the domain data files, keeping the LG 100% aligned
with the slide deck, Lesson Plan and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5; from data_domain6 import DOMAIN6
from data_domain7 import DOMAIN7
from data_files import DATA
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5+DOMAIN6+DOMAIN7
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"activities")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def img(path,caption): B.append(("img",path,caption))
def rule(): B.append(("rule",))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  "It provides detailed step-by-step instructions for all 12 hands-on activities, organised by the "
  f"seven course topics aligned to the Skills Framework TSC {C.TSC_TITLE} ({C.TSC_CODE}). "
  "The activities use Xero, Microsoft Power BI, Google Forms and case-study templates so that every "
  "learning outcome is practised hands-on.")
p("Use this guide alongside the course slides and the activity folders (activities/) of the course "
  "repository. The slides give each activity a one-page overview; this guide carries the full detailed "
  "steps. The final assessment is open book — you may refer to the slides, this Learner Guide and any "
  "approved materials.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Skills Framework Alignment")
p(f"This course is aligned to the {C.TSC_TITLE} TSC ({C.TSC_CODE}) under the ICT Skills Framework.")
h3("TSC Abilities")
bullets(C.TSC_ABILITIES)
h3("TSC Knowledge")
bullets(C.TSC_KNOWLEDGE)

h1("Before You Start — Environment Setup")
h3("What you need")
bullets([
 "A laptop with a modern browser (Chrome, Edge or Safari) and internet access.",
 "An email address you can access in class — used to sign up for the free Xero trial and Microsoft Power BI accounts.",
 "The lab data files (Excel workbooks + CSV mirrors) — one folder per activity, activities/activity01 … activity12, in the course repository and on the LMS (lms-tms.tertiaryinfotech.com) under Activities. Every activity's mock data (financial statements, budgets, actuals, the 99 Agency template, Xero-style exports and tax worksheets) ships there.",
 "Excel, Google Sheets or LibreOffice Calc to work through the case-study schedules.",
])
h3("Accounts you will create during the course")
bullets([
 "Xero (free 30-day trial, includes the Demo Company) — https://www.xero.com/sg/signup/",
 "Microsoft Power BI (free trial) — https://powerbi.microsoft.com/en-us/landing/signin/",
])
h3("Conventions used in every activity")
bullets([
 "Menu paths are written as Menu → Submenu → Item (e.g. Accounting → Reports → Budget Manager).",
 "Amounts are in Singapore dollars unless stated otherwise.",
 "Each activity ends with a 'Check your work' section — verify before moving on.",
 "If a screen differs slightly from these steps, Xero or Power BI may have updated its UI; the flow remains the same.",
])

# ---------------- per-topic, per-lab ----------------
for t in C.TOPICS:
    h1(f"Topic {t['code']} — {t['title']}")
    p(t["subtitle"])
    h3("Key concepts")
    bullets([f"{k} — {v}" for k,v in t["concepts"]])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        h2(f"Activity {a['num']} — {a['title']}")
        p(f"Learning outcome: {a['objective']}.")
        p(f"Goal: {a['desc']}")
        h3("What you'll produce")
        p(a["build"]+f"   (Tools: {a['services']}.)")
        img(os.path.join(REPO,"activities",f"activity{a['num']:02d}",f"activity-{a['num']:02d}-workflow.png"),
            f"Activity {a['num']} workflow — {a['title']}")
        h3("Step-by-step")
        st=[]
        for i,(instr,cmd) in enumerate(a["steps"],1):
            st.append((instr,cmd))
        steps(st)
        d=DATA.get(a["num"])
        if d:
            h3(f"Data files for this activity (activities/activity{a['num']:02d}/)")
            bullets([f"{fn} — {desc}" for fn,desc in d["files"]])
            h3("Analyzing the Excel workbook — step by step")
            steps([(s,"") for s in d["excel_steps"]])
            h3("Analyzing the CSV data — step by step")
            steps([(s,"") for s in d["csv_steps"]])
        h3("Check your work")
        p(a["test"])
        note(f"A printable copy of this activity — with its workflow diagram and data files — is in the activities/activity{a['num']:02d}/ folder of the course repository.")
        rule()

h1("Assessment Preparation")
bullets([
 "The Written Assessment (SAQ) is 70 minutes and tests the knowledge areas K1–K10 — review the Key Concepts of every topic.",
 "The Written Assessment (Case Study) is 80 minutes and tests the abilities A1–A7 — rehearse the 99 Agency master budget and the variance analysis activities.",
 "Both papers are open book: bring the slides, this Learner Guide and your completed activity templates.",
 "A minimum of 75% attendance is required to be eligible for assessment and funding.",
])

h1("Glossary")
gl=[
 ("Budget","A quantitative plan for acquiring and using resources over a specified period."),
 ("Master Budget","The document aggregating all departmental inputs — budgeted P&L, Balance Sheet and Cash Flow."),
 ("Baseline / Incremental / Zero-based","Budget preparation methods: previous plan as base / % or $ on the baseline / starting fresh."),
 ("Accrual accounting","Recording revenue when earned and expenses when incurred, regardless of cash movement."),
 ("Time value of money","A dollar today is worth more than a dollar tomorrow, due to its earning capacity."),
 ("Cash budget","An estimate of cash inflows and outflows over a period, proving the entity can keep operating."),
 ("Working capital","Current assets minus current liabilities — the measure of short-term financial health."),
 ("Budgetary control","Comparing budget standards with actual performance and taking corrective action."),
 ("Variance","Actual minus budget, in dollars or %; favourable (better than budget) or adverse (worse)."),
 ("Vertical / Horizontal analysis","Each line as a % of a base within a period / change of each line across periods."),
 ("Year of Assessment (YA)","The 12-month period in which a company's income is assessed by IRAS."),
 ("ECI","Estimated Chargeable Income — filed within 3 months of the financial year end."),
 ("Start-Up Tax Exemption","First 3 YAs: 75% exemption on the first S$100k and 50% on the next S$100k of chargeable income."),
]
B.append(("dl",gl))

# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="code": out+=["```",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="img":
            rel=os.path.relpath(rest[0],REPO)
            out+=[f"![{rest[1]}]({rel})",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("10",C.VERSION_DATE,"Legacy release — Budgeting for SMEs courseware (reference deck v10).",C.TRAINER),
 ("11",C.VERSION_DATE,"Full regeneration in the single-source pipeline: detailed step-by-step guides for all 12 hands-on activities across the 7 topics, aligned to the visual slide deck and the WA (SAQ + CS) assessments.",C.TRAINER),
 ("12",C.VERSION_DATE,"QA fixes: workflow diagram embedded per activity, per-activity step numbering matching the Markdown mirror.",C.TRAINER),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,"Added mock data sets for every activity (Excel workbooks + CSV mirrors, one folder per activity: activities/activity01 - activity12) with full step-by-step walkthroughs for analyzing both the Excel and the CSV data.",C.TRAINER),
])
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph()
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        for i,(instr,cmd) in enumerate(rest[0],1):
            para=doc.add_paragraph()
            para.paragraph_format.left_indent=Pt(18)
            para.paragraph_format.first_line_indent=Pt(-18)
            r=para.add_run(f"{i}. "); r.bold=True; r.font.color.rgb=BRAND
            para.add_run(instr)
            if cmd: code_para(cmd)
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="img":
        from docx.shared import Inches as _In
        if os.path.exists(rest[0]):
            doc.add_picture(rest[0],width=_In(6.4))
            cap=doc.add_paragraph(); cr=cap.add_run(rest[1]); cr.italic=True; cr.font.size=Pt(9); cr.font.color.rgb=GREY
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
